from bs4 import BeautifulSoup, Tag, NavigableString
import requests
import urllib3
import xlsxwriter
from docx import Document

cand_texts = []
alts = []
names = []
universal_count = 0

def cand_info_extract(cand_url, img_url, alt, name):
    source = requests.get(cand_url).content
    cand_soup = BeautifulSoup(source, 'html.parser')

    image = requests.get(img_url)
    image_path = '<your_computer_directory_path>/Desktop/cand_images/' + str(name) + ".png"
    open(image_path, 'wb').write(image.content)
    
    #navigating to relevant section of the page
    tables = cand_soup.find_all('table')
    req_table = tables[14]

    rows = req_table.find_all('tr')
    rows.pop()
    rows.pop()
    nrows = len(rows)

    #getting education and experience
    h3s = cand_soup.find_all('h3')
    req_h3 = h3s[-3]
    education = ""
    next_s = req_h3.nextSibling

    while next_s:
        if next_s in h3s:
            break

        if isinstance(next_s, NavigableString):
            education += next_s
            print(next_s)
            next_s = next_s.nextSibling
        elif next_s in cand_soup.find_all('hr'):
            education += "\n"
            next_s = next_s.nextSibling
        elif next_s in cand_soup.find_all('br'):
            education += "\n"
            next_s = next_s.nextSibling
        elif next_s in cand_soup.find_all('b'):
            education += next_s.get_text()
            next_s = next_s.nextSibling
        elif isinstance(next_s, Tag):
            next_s = next_s.nextSibling

    req_h3 = h3s[-4]
    exp = ""
    next_s = req_h3.nextSibling

    while next_s:
        if next_s in h3s:
            break

        if isinstance(next_s, NavigableString):
            exp += next_s
            print(next_s)
            next_s = next_s.nextSibling
        elif next_s in cand_soup.find_all('hr'):
            exp += "\n"
            next_s = next_s.nextSibling
        elif next_s in cand_soup.find_all('br'):
            exp += "\n"
            next_s=next_s.nextSibling
        elif next_s in cand_soup.find_all('b'):
            exp += next_s.get_text()
            next_s = next_s.nextSibling
        elif isinstance(next_s, Tag):
            next_s = next_s.nextSibling

    if len(education) == 0:
        education = "--"
    if len(exp) == 0:
        exp = "--"
        
    #education and exp ends        

    headers = []
    values = []

    for i in range(nrows):
        info = rows[i].find_all('td')
        headers.append(info[0].get_text())
        values.append(info[1].get_text())

    headers.append("\nExperience : \n")
    headers.append("\nEducation : \n")
    values.append(exp)
    values.append(education)

    #making text for the candidate which will be entered into the output file
    text = "Name : " + name + "\n"
    for j in range(nrows + 2):
        text += headers[j] + values[j] + "\n"
    
    cand_texts.append(text)
    print(name)

def page_parser(url):

    source = requests.get(url).content
    soup = BeautifulSoup(source, 'html.parser')

    #navigating to the relevant page section
    cand_div = soup.find('div', attrs={'id': 'ctl00_ctl00_ContentPlaceHolder1_ContentPlaceHolder1_divContent'})
    main_table = cand_div.find('table', attrs={'width': '100%'})

    #getting all required rows in the page
    rows = main_table.find_all('tr')
    for i in range(2):
        rows.pop()

         
    count = 0
    for row in rows:
        data = row.find('td')
        if len(data.find_all('script')) or len(data.find_all('hr')):
            rows.pop(count)

        count += 1

    nrows = len(rows)
    print(nrows)

    #extracting information of each candidate on page
    for i in range(2, nrows):
        cand_row = rows[i]
        cand_data = cand_row.find_all('td')

        img_url = "https://www.myvisajobs.com" + cand_data[0].find('img').attrs['src']
        alt = cand_data[0].find('img').attrs['alt']
        alts.append(alt)
        cand_url = "https://www.myvisajobs.com" + cand_data[0].find('a').attrs['href']

        name = cand_data[1].find('a').get_text()
        names.append(name)

        #extracting info from candidate data
        cand_info_extract(cand_url, img_url, alt, name)

    return True
  

#function to write data to an excel file
def output_file(cand_texts, names):

    workbook = xlsxwriter.Workbook('output.xlsx')
    worksheet = workbook.add_worksheet()

    for i in range(len(cand_texts)):
        worksheet.write(i, 0, str(i+1))
        worksheet.write(i, 1, cand_texts[i])

    workbook.close()

    doc = Document()
    i = 0
    for i in range(len(cand_texts)):
        img_path = "<your_computer_directory_path>/Desktop/cand_images/" + names[i] + ".png"

        para = doc.add_paragraph()
        run = para.add_run()
        text = str(i+1) + ". " 
        run.add_text(text)
        run.add_picture(img_path)

    doc.save("image_doc.docx")



# Main functioning of the program
#for m in range(1, 16):
m = 1

if m == 1:
    url = "https://www.myvisajobs.com/CV/Candidates.aspx"
else:
    url = "https://www.myvisajobs.com/CV/Candidates.aspx" + "?P=" + str(m)

page_parser(url)

output_file(cand_texts, names)


